#!/usr/bin/env python3
"""Generate a branded Flowistic proposal document for BeyondBookings."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand colors
TEAL = RGBColor(0x0A, 0xA8, 0xA7)
DARK_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
TEAL_LIGHT = RGBColor(0xE6, 0xF7, 0xF7)
TEAL_BG = "E6F7F7"
TEAL_HEX = "0AA8A7"
DARK_HEX = "4A4A4A"
WHITE_HEX = "FFFFFF"
LIGHT_HEX = "F5F5F5"

BRAND_DIR = "/Users/fayssalelmofatiche/dev/ideas/brand"
OUTPUT_DIR = "/Users/fayssalelmofatiche/dev/ideas/projects/beyondbookings"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = DARK_GRAY

# Heading styles
for i, (size, color) in enumerate([
    (Pt(26), TEAL),   # Heading 1
    (Pt(16), TEAL),   # Heading 2
    (Pt(13), DARK_GRAY),  # Heading 3
    (Pt(11), TEAL),   # Heading 4
], start=1):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if i <= 2 else 12)
    h.paragraph_format.space_after = Pt(6)


def add_teal_line():
    """Add a teal horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TEAL_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size or Pt(9.5)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)


def add_styled_table(headers, rows, col_widths=None, highlight_last=False):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Remove default borders, we'll style manually
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TEAL_HEX)
        set_cell_text(cell, header, bold=True, color=WHITE, size=Pt(9.5))

    # Data rows
    for r, row_data in enumerate(rows):
        is_last = (r == len(rows) - 1) and highlight_last
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            if is_last:
                set_cell_shading(cell, TEAL_BG)
                set_cell_text(cell, str(value), bold=True, color=TEAL, size=Pt(9.5))
            elif r % 2 == 1:
                set_cell_shading(cell, LIGHT_HEX)
                set_cell_text(cell, str(value), size=Pt(9.5))
            else:
                set_cell_text(cell, str(value), size=Pt(9.5))

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    return table


def add_bullet(text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = DARK_GRAY
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
    else:
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK_GRAY
        if not p.runs:
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK_GRAY


def add_body(text, bold=False, italic=False):
    """Add body text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Dual logos using a borderless table (Flowistic left, BeyondBookings right)
logo_table = doc.add_table(rows=1, cols=2)
logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Remove all borders
for row in logo_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# Flowistic logo (left)
flowistic_logo = os.path.join(BRAND_DIR, "logo_horizontal_bright_bg.png")
cell_left = logo_table.rows[0].cells[0]
cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
run = cell_left.paragraphs[0].add_run()
run.add_picture(flowistic_logo, width=Cm(7))

# BeyondBookings logo (right)
bb_logo = os.path.join(BRAND_DIR, "beyond_booking_logo.png")
cell_right = logo_table.rows[0].cells[1]
cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = cell_right.paragraphs[0].add_run()
run.add_picture(bb_logo, width=Cm(5))

# Spacer
for _ in range(3):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Invoice Payment Automation")
run.font.name = 'Calibri'
run.font.size = Pt(32)
run.font.color.rgb = TEAL
run.font.bold = True

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Effort & Cost Estimation")
run.font.name = 'Calibri'
run.font.size = Pt(20)
run.font.color.rgb = DARK_GRAY

# Spacer
doc.add_paragraph()

# Client info
add_teal_line()

# Prepared for
p = doc.add_paragraph()
run = p.add_run("Prepared for: ")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY
run = p.add_run("Dr. Stefan Kloss")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("                        ")
run.font.size = Pt(12)
run = p.add_run("CFO, BeyondBookings")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY

# Spacer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

# Prepared by
p = doc.add_paragraph()
run = p.add_run("Prepared by:  ")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY
run = p.add_run("Fayssal El Mofatiche")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("                        ")
run.font.size = Pt(12)
run = p.add_run("CEO & Founder, Flowistic")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY

# Spacer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run("Date: ")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY
run = p.add_run("February 2025")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY

p = doc.add_paragraph()
run = p.add_run("Version: ")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY
run = p.add_run("1.0 (Draft)")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = DARK_GRAY

# Confidentiality
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("CONFIDENTIAL")
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = TEAL
run.font.bold = True
p = doc.add_paragraph()
run = p.add_run("This document contains proprietary information prepared exclusively for BeyondBookings. "
                "Distribution without written consent from Flowistic is not permitted.")
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = DARK_GRAY
run.font.italic = True

# Page break
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════

doc.add_heading("Executive Summary", level=1)
add_teal_line()

add_body(
    "BeyondBookings currently processes approximately 1,000 supplier invoices per year through a "
    "dedicated Outlook mailbox. Each invoice is manually reviewed, data is extracted by hand, and "
    "payments are initiated via manual bank transfers. Additionally, the CFO maintains a local "
    "Access database to consolidate financial data, manually pulling customer information from HubSpot."
)

add_body(
    "This proposal outlines an AI-powered automation solution that eliminates these manual workflows. "
    "The solution is structured as modular components that can be delivered incrementally, allowing "
    "BeyondBookings to realize value early while managing risk and investment."
)

doc.add_heading("Proposed Components", level=3)

add_bullet("AI-powered invoice extraction and approval workflow (always included)", bold_prefix="Core: ")
add_bullet("Automated payment execution and operational hardening", bold_prefix="Variant A: ")
add_bullet("Cloud-based CFO financial consolidation replacing the Access database", bold_prefix="Variant B: ")
add_bullet("Client-facing invoice submission portal replacing the Outlook inbox", bold_prefix="Variant C: ")

doc.add_heading("Investment Overview", level=3)

add_styled_table(
    headers=["Package", "Scope", "Effort", "Cost"],
    rows=[
        ["A: Core only", "Core", "19-26 PD", "\u20ac22,800 - \u20ac31,200"],
        ["B: Core + CFO", "Core + Var. B", "47-64 PD", "\u20ac49,350 - \u20ac67,200"],
        ["C: Core + Portal", "Core + Var. C", "47-64 PD", "\u20ac49,350 - \u20ac67,200"],
        ["D: Core + CFO + Portal", "Core + Var. B + C", "75-102 PD", "\u20ac71,250 - \u20ac96,900"],
        ["E: Everything", "Core + Var. A + B + C", "93-126 PD", "\u20ac81,840 - \u20ac110,880"],
    ],
    col_widths=[4.5, 5, 3, 4.5],
    highlight_last=False,
)

add_body("")
add_body(
    "We recommend Package D as the optimal engagement: it delivers Core automation plus both "
    "the CFO dashboard and client portal at a 20% reduced rate, while deferring "
    "payment automation (Variant A) until invoice extraction accuracy is validated.",
    bold=True,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# UNDERSTANDING
# ══════════════════════════════════════════════════════════════

doc.add_heading("Understanding Your Needs", level=1)
add_teal_line()

doc.add_heading("Current Situation", level=2)

add_bullet("~1,000 invoices/year received via a dedicated Outlook mailbox (up from 500 last year)")
add_bullet("Mix of digitally generated and scanned PDF invoices")
add_bullet("Manual data extraction, review, and payment initiation via bank transfer")
add_bullet("Single approver workflow (lean team of 10-12 people)")
add_bullet("CFO manually consolidates financials in a local Access database")
add_bullet("Customer data is pulled manually from HubSpot into the Access database")

doc.add_heading("Current Technology Stack", level=2)

add_styled_table(
    headers=["System", "Purpose", "Status"],
    rows=[
        ["Outlook", "Invoice reception (dedicated mailbox)", "To be automated / replaced"],
        ["Lexoffice", "Bookkeeping", "Keep \u2013 integrate via API"],
        ["HubSpot", "CRM \u2013 customer & deal management", "Keep \u2013 integrate via API"],
        ["Access Database", "CFO financial consolidation", "To be replaced"],
        ["Manual Bank Transfer", "Payment execution", "To be automated (Variant A)"],
    ],
    col_widths=[3.5, 6, 4.5],
)

doc.add_heading("Key Challenges", level=2)

add_bullet("Time-consuming manual invoice processing prone to human error")
add_bullet("No real-time visibility into invoice status for suppliers")
add_bullet("Fragmented financial data across multiple disconnected systems")
add_bullet("Local Access database creates single-point-of-failure risk and limits collaboration")
add_bullet("Scaling to higher invoice volumes will compound manual effort")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CORE - INVOICE PROCESSING MVP
# ══════════════════════════════════════════════════════════════

doc.add_heading("Core: Invoice Processing MVP", level=1)
add_teal_line()

add_body(
    "The core of the solution: an AI agent that monitors the Outlook mailbox, extracts invoice data "
    "using OCR and language models, and presents it for approval in a dashboard. Approved invoices "
    "are pushed directly to Lexoffice."
)

doc.add_heading("Detailed Breakdown", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["Email Ingestion", "Microsoft Graph API integration to monitor the Outlook mailbox, extract PDF attachments, queue for processing", "2-3", "16-24h"],
        ["PDF Processing", "Digital PDF text extraction + OCR for scanned documents (Azure Document Intelligence)", "3-4", "24-32h"],
        ["AI Data Extraction", "LLM-based field extraction (vendor, invoice #, date, due date, amount, VAT, IBAN, line items) with confidence scoring", "4-5", "32-40h"],
        ["Approval Dashboard", "Web UI: pending invoices list, side-by-side PDF viewer + extracted data, approve/reject/edit, email notifications", "3-4", "24-32h"],
        ["Lexoffice Integration", "Push approved invoice data into Lexoffice via API", "2-3", "16-24h"],
        ["Infrastructure & Auth", "Hosting, authentication, basic monitoring, environment setup", "2-3", "16-24h"],
        ["Testing & QA", "End-to-end testing with real invoices, edge case handling, accuracy tuning", "3-4", "24-32h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Core Total: 19-26 PD (152-208 hours)")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = TEAL
run.font.bold = True

doc.add_heading("Deliverables", level=3)

add_bullet("Invoices automatically read from Outlook and parsed by AI")
add_bullet("Extracted data presented for review in an approval dashboard")
add_bullet("Approver can review, correct, and approve or reject each invoice")
add_bullet("Approved invoices pushed to Lexoffice with all fields pre-filled")
add_bullet("Payment still triggered manually (fully prepared in Lexoffice)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# VARIANT A - FULL PAYMENT AUTOMATION
# ══════════════════════════════════════════════════════════════

doc.add_heading("Variant A: Full Payment Automation", level=1)
add_teal_line()

add_body(
    "Once invoice extraction accuracy is validated through Core, this variant adds automated "
    "payment execution and operational improvements based on real-world learnings."
)

doc.add_heading("Detailed Breakdown", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["Payment Automation", "Banking API integration (e.g. finAPI) or Lexoffice payment triggering for approved invoices", "5-6", "40-48h"],
        ["Duplicate Detection", "Invoice deduplication, vendor matching against known database", "3-4", "24-32h"],
        ["Accuracy Feedback Loop", "Use approver corrections to improve extraction accuracy over time", "3-4", "24-32h"],
        ["Notifications & Alerts", "Reminders for pending approvals, payment confirmations, failure alerts", "2-3", "16-24h"],
        ["Reporting", "Dashboard: processed/pending/rejected counts, monthly spend overview", "2-3", "16-24h"],
        ["Hardening", "Error handling, retry logic, GDPR data handling review, audit logging", "3-4", "24-32h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Variant A Total: 18-24 PD (144-192 hours)")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = TEAL
run.font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# VARIANT B - CFO CONSOLIDATION
# ══════════════════════════════════════════════════════════════

doc.add_heading("Variant B: CFO Financial Consolidation", level=1)
add_teal_line()

add_body(
    "Replace the local Access database with a cloud-based financial consolidation layer. "
    "Automates data feeds from HubSpot and Lexoffice, and connects to the invoice automation "
    "pipeline from Core. Can be executed in parallel with Core."
)

doc.add_heading("Discovery & Data Migration (1-2 weeks)", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["Access DB Analysis", "Reverse-engineer the existing database: tables, relationships, queries, reports", "3-4", "24-32h"],
        ["Data Model Design", "Design a clean, normalized data model for financial consolidation needs", "2-3", "16-24h"],
        ["Data Migration", "Migrate historical data from Access. Validate completeness with the CFO", "2-3", "16-24h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

doc.add_heading("Automated Data Integrations (2-3 weeks)", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["HubSpot Integration", "Automated sync of customer data, deals, and revenue via HubSpot API. Scheduled sync + webhooks", "4-5", "32-40h"],
        ["Lexoffice Integration", "Pull bookkeeping data (expenses, revenue, tax) from Lexoffice API", "3-4", "24-32h"],
        ["Invoice Pipeline Link", "Connect Core invoice automation pipeline so approved/paid invoices flow into CFO view", "2-3", "16-24h"],
        ["Data Reconciliation", "Automated cross-checks between sources with discrepancy flagging", "3-4", "24-32h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

doc.add_heading("CFO Dashboard (2-3 weeks)", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["Financial Overview", "Consolidated dashboard: revenue, expenses, cash flow, outstanding invoices, AP/AR", "4-5", "32-40h"],
        ["Customer Financial View", "Per-customer P&L combining HubSpot deal data with invoice/payment history", "3-4", "24-32h"],
        ["Reporting & Export", "Monthly/quarterly reports, CSV/PDF export for accountant or tax advisor", "2-3", "16-24h"],
        ["Filters & Drill-Down", "Filter by period, customer, vendor, category. Drill down to transactions", "2-3", "16-24h"],
        ["Testing & Validation", "Validate numbers against Access database, iterate based on CFO feedback", "3-4", "24-32h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Variant B Total: 28-38 PD (224-304 hours)")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = TEAL
run.font.bold = True

doc.add_heading("Deliverables", level=3)

add_bullet("Local Access database fully replaced by cloud-based solution")
add_bullet("HubSpot customer/revenue data synced automatically")
add_bullet("Lexoffice bookkeeping data integrated automatically")
add_bullet("Invoice pipeline data flows in from Core automation")
add_bullet("CFO dashboard with consolidated financial overview, per-customer views, and exportable reports")
add_bullet("Discrepancy detection between data sources")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# VARIANT C - CLIENT PORTAL
# ══════════════════════════════════════════════════════════════

doc.add_heading("Variant C: Client Invoice Portal", level=1)
add_teal_line()

add_body(
    "Replace the Outlook mailbox with a branded client-facing portal where suppliers upload "
    "invoices directly, track processing status in real-time, and communicate with BeyondBookings. "
    "Improves client experience and data quality through structured submissions."
)

doc.add_heading("Portal Core (3-4 weeks)", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["Client Authentication", "Secure login with invitation-based onboarding. Role management for multiple users per supplier", "4-5", "32-40h"],
        ["Invoice Upload", "Drag-and-drop interface. Single/bulk PDF upload. Optional structured fields to improve extraction", "4-5", "32-40h"],
        ["Status Tracking", "Real-time status: Uploaded \u2192 Processing \u2192 Under Review \u2192 Approved \u2192 Paid. Timeline per invoice", "4-5", "32-40h"],
        ["Feedback & Communication", "Per-invoice comment thread. Rejection reasons with resubmission workflow", "3-4", "24-32h"],
        ["Email Notifications", "Automated notifications on status changes. Configurable preferences", "2-3", "16-24h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

doc.add_heading("Integration & Migration (1-2 weeks)", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["Pipeline Integration", "Connect portal uploads to Core processing pipeline. Dual-mode during transition", "3-4", "24-32h"],
        ["Dashboard Updates", "Extend approval dashboard with client feedback capabilities", "2-3", "16-24h"],
        ["Client Onboarding Flow", "Admin interface to invite clients, manage accounts, monitor usage", "2-3", "16-24h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

doc.add_heading("Branding & Polish (1 week)", level=2)

add_styled_table(
    headers=["Task", "Description", "PD", "Hours"],
    rows=[
        ["UI/UX Design", "Branded interface aligned with BeyondBookings identity. Mobile-responsive", "3-4", "24-32h"],
        ["Multi-language", "German and English support", "1-2", "8-16h"],
        ["Testing & QA", "End-to-end testing, cross-browser, mobile testing", "3-4", "24-32h"],
    ],
    col_widths=[3.5, 7, 1.5, 2],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Variant C Total: 28-38 PD (224-304 hours)")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = TEAL
run.font.bold = True

doc.add_heading("Deliverables", level=3)

add_bullet("Branded client portal with secure login for suppliers")
add_bullet("Drag-and-drop invoice upload (single and bulk)")
add_bullet("Real-time status tracking for every submitted invoice")
add_bullet("Per-invoice feedback and communication thread")
add_bullet("Automated email notifications on status changes")
add_bullet("Dual-mode ingestion: portal + Outlook during transition period")
add_bullet("Admin interface for client account management")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# COST ESTIMATION
# ══════════════════════════════════════════════════════════════

doc.add_heading("Investment", level=1)
add_teal_line()

doc.add_heading("Rate Structure", level=2)

add_body(
    "Rates are tiered based on total engagement volume, rewarding larger commitments "
    "with a reduced daily rate:"
)

add_styled_table(
    headers=["Engagement Volume", "Daily Rate", "Hourly Rate"],
    rows=[
        ["Up to 40 PD", "\u20ac1,200/day", "\u20ac150/h"],
        ["41-80 PD", "\u20ac1,050/day", "\u20ac130/h"],
        ["81-120 PD", "\u20ac950/day", "\u20ac120/h"],
        ["120+ PD", "\u20ac880/day", "\u20ac110/h"],
    ],
    col_widths=[5, 4.5, 4.5],
)

doc.add_heading("Per-Component Effort", level=2)

add_styled_table(
    headers=["Component", "Effort (PD)", "Hours"],
    rows=[
        ["Core (Invoice Processing MVP)", "19-26 PD", "152-208h"],
        ["Variant A (Full Payment Automation)", "18-24 PD", "144-192h"],
        ["Variant B (CFO Consolidation)", "28-38 PD", "224-304h"],
        ["Variant C (Client Portal)", "28-38 PD", "224-304h"],
    ],
    col_widths=[6, 4, 4],
)

doc.add_heading("Packaging Options", level=2)

add_styled_table(
    headers=["Package", "Scope", "Effort", "Rate", "Investment"],
    rows=[
        ["A: Core only", "Core", "19-26 PD", "\u20ac1,200/day", "\u20ac22,800 - \u20ac31,200"],
        ["B: Core + CFO", "Core + Var. B", "47-64 PD", "\u20ac1,050/day", "\u20ac49,350 - \u20ac67,200"],
        ["C: Core + Portal", "Core + Var. C", "47-64 PD", "\u20ac1,050/day", "\u20ac49,350 - \u20ac67,200"],
        ["D: Core + CFO + Portal", "Core + Var. B + C", "75-102 PD", "\u20ac950/day", "\u20ac71,250 - \u20ac96,900"],
        ["E: Everything", "Core + Var. A + B + C", "93-126 PD", "\u20ac880/day", "\u20ac81,840 - \u20ac110,880"],
    ],
    col_widths=[3.5, 3.5, 2.5, 2.5, 4],
)

doc.add_heading("Savings at Scale", level=2)

add_styled_table(
    headers=["Package", "Rate Reduction", "Saving / PD", "Total Saving"],
    rows=[
        ["A", "Baseline", "-", "-"],
        ["B / C", "-12.5%", "\u20ac150/day", "\u20ac7,050 - \u20ac9,600"],
        ["D", "-20.8%", "\u20ac250/day", "\u20ac18,750 - \u20ac25,500"],
        ["E", "-26.7%", "\u20ac320/day", "\u20ac29,760 - \u20ac40,320"],
    ],
    col_widths=[3.5, 3.5, 3.5, 5.5],
)

doc.add_heading("Ongoing Infrastructure Costs", level=2)

add_body("Estimated monthly running costs once the solution is in production:")

add_styled_table(
    headers=["Service", "Estimated Monthly Cost"],
    rows=[
        ["OCR Service (provider TBD)", "\u20ac50-100"],
        ["LLM API (GPT-4o / Claude)", "\u20ac30-80"],
        ["Cloud Hosting", "\u20ac30-60"],
        ["Database (PostgreSQL)", "\u20ac20-50"],
        ["Email API (Microsoft Graph)", "Included in M365"],
        ["Lexoffice API", "Included in subscription"],
        ["Banking API (Variant A)", "\u20ac50-150"],
        ["Portal Auth & Email (Variant C)", "\u20ac0-45"],
    ],
    col_widths=[8, 6],
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Estimated total infrastructure: \u20ac200-550/month (all components)")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = TEAL
run.font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# RISKS
# ══════════════════════════════════════════════════════════════

doc.add_heading("Risks & Considerations", level=1)
add_teal_line()

add_styled_table(
    headers=["Risk", "Impact", "Mitigation"],
    rows=[
        ["OCR accuracy on scanned invoices", "High", "Budget tuning time, confidence scoring flags uncertain extractions for manual review"],
        ["Diverse invoice layouts", "Medium", "LLM-based extraction handles variety well; feedback loop improves accuracy over time"],
        ["Banking API complexity", "Medium", "Start with Lexoffice-triggered payments before full bank API integration"],
        ["GDPR / data handling", "Medium", "Flag for legal review; ensure minimal data retention"],
        ["GoBD compliance", "Low-Med", "Client to verify with tax advisor; may need compliant archiving"],
        ["Access DB data quality", "Medium", "Legacy data may be inconsistent; budget time for cleanup and CFO validation"],
        ["HubSpot data model", "Low-Med", "Map custom properties to financial model in discovery session"],
        ["CFO adoption", "Medium", "Involve CFO early in dashboard design; iterative feedback sessions"],
        ["Client portal adoption", "Medium", "Clear UX advantage and status tracking incentivize suppliers to switch from email"],
    ],
    col_widths=[4, 2, 8],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# RECOMMENDATION & TIMELINE
# ══════════════════════════════════════════════════════════════

doc.add_heading("Recommendation", level=1)
add_teal_line()

add_body(
    "We recommend Package D (Core + Variant B + Variant C) as the optimal engagement. "
    "This delivers Core automation plus the CFO dashboard and client portal at a 20% reduced rate, while "
    "deferring payment automation (Variant A) until the system is proven.",
    bold=True,
)

doc.add_heading("Proposed Timeline", level=2)

add_styled_table(
    headers=["Period", "Core (MVP)", "Variant B (CFO)", "Variant C (Portal)"],
    rows=[
        ["Weeks 1-4", "Core pipeline development", "Access DB analysis + data model", "UI/UX design (parallel)"],
        ["Weeks 5-8", "Dashboard + integrations", "HubSpot & Lexoffice integration", "Portal core development"],
        ["Weeks 9-10", "Testing & go-live", "CFO dashboard + validation", "Integration + testing"],
        ["Weeks 11-12", "Support & tuning", "Refinement", "Client onboarding"],
    ],
    col_widths=[2.5, 4, 4, 4],
)

doc.add_heading("Value Delivered", level=2)

add_bullet("Automated extraction and approval workflow, eliminating manual data entry", bold_prefix="Invoice processor: ")
add_bullet("Real-time financial consolidation replacing manual Access work", bold_prefix="CFO: ")
add_bullet("Self-service portal with full transparency on invoice status", bold_prefix="Suppliers: ")

doc.add_heading("Alternative: Start Smaller", level=2)

add_body(
    "If budget is a constraint, Package B (Core + Variant B) addresses the two most "
    "painful manual workflows at \u20ac49,350-\u20ac67,200. The client portal can follow as a separate engagement."
)

# ══════════════════════════════════════════════════════════════
# NEXT STEPS
# ══════════════════════════════════════════════════════════════

doc.add_heading("Next Steps", level=1)
add_teal_line()

add_bullet("Review this proposal and confirm preferred package")
add_bullet("Schedule a discovery session to review the Access database and HubSpot data model")
add_bullet("Provide sample invoices (5-10 representative PDFs) for extraction accuracy testing")
add_bullet("Confirm Lexoffice API access and Microsoft 365 admin permissions")
add_bullet("Align on timeline and kick-off date")

doc.add_paragraph()
add_teal_line()

# Closing logo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run()
run.add_picture(os.path.join(BRAND_DIR, "graphic_only.png"), width=Cm(2.5))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Fayssal El Mofatiche")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CEO & Founder, Flowistic")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = DARK_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Dare to deliver.")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("flowistic.ai")
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = TEAL

# ── Save ──
output_path = os.path.join(OUTPUT_DIR, "BeyondBookings_Proposal_Flowistic.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
