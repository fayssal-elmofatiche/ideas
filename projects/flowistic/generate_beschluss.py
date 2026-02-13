#!/usr/bin/env python3
"""Generate a branded Flowistic Gesellschafterbeschluss for salary adjustment."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Brand colors
TEAL = RGBColor(0x0A, 0xA8, 0xA7)
DARK_GRAY = RGBColor(0x4A, 0x4A, 0x4A)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x33, 0x33, 0x33)
TEAL_HEX = "0AA8A7"
LIGHT_HEX = "F5F5F5"

BRAND_DIR = "/Users/fayssalelmofatiche/dev/ideas/brand"
OUTPUT_DIR = "/Users/fayssalelmofatiche/dev/ideas/projects/flowistic"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY


def add_teal_line():
    """Add a teal horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TEAL_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_body(text, bold=False, italic=False, size=Pt(11), align=None):
    """Add body text."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size
    run.font.color.rgb = DARK_GRAY
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size or Pt(10)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)


# ══════════════════════════════════════════════════════════════
# HEADER WITH LOGO
# ══════════════════════════════════════════════════════════════

logo_path = os.path.join(BRAND_DIR, "logo_horizontal_bright_bg.png")
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(logo_path, width=Cm(5))
    p.paragraph_format.space_after = Pt(4)

add_teal_line()

# ══════════════════════════════════════════════════════════════
# DOCUMENT TITLE
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Gesellschafterbeschluss")
run.font.name = 'Calibri'
run.font.size = Pt(24)
run.font.color.rgb = TEAL
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run("Anpassung der Geschäftsführervergütung")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = DARK_GRAY

add_teal_line()

# ══════════════════════════════════════════════════════════════
# COMPANY DETAILS
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Gesellschaft")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

# Company info table (borderless)
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

info_data = [
    ("Firma:", "Flowistic GmbH"),
    ("Sitz:", "Kostheimer Str. 14, 60326 Frankfurt am Main"),
    ("Registergericht:", "Amtsgericht Frankfurt am Main"),
    ("Handelsregisternummer:", "HRB 131160"),
]

for row_idx, (label, value) in enumerate(info_data):
    # Remove borders
    for cell in info_table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    set_cell_text(info_table.rows[row_idx].cells[0], label, bold=True, size=Pt(10.5))
    set_cell_text(info_table.rows[row_idx].cells[1], value, size=Pt(10.5))
    info_table.rows[row_idx].cells[0].width = Cm(5)
    info_table.rows[row_idx].cells[1].width = Cm(10)

# ══════════════════════════════════════════════════════════════
# BESCHLUSS
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Beschlussfassung")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

add_body(
    "Der alleinige Gesellschafter der Flowistic GmbH, Herr Fayssal El Mofatiche, "
    "fasst hiermit folgenden Beschluss im Umlaufverfahren gemäß § 48 Abs. 2 GmbHG:"
)

# Beschluss box with teal left border
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)
p.paragraph_format.left_indent = Cm(1)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{TEAL_HEX}"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

run = p.add_run("Beschluss Nr. 2026-01")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = TEAL
run.font.bold = True

run = p.add_run("\n\n")

run = p.add_run(
    "Das monatliche Bruttogehalt des Geschäftsführers Fayssal El Mofatiche wird "
    "mit Wirkung zum 01. Februar 2026 von 5.000,00 € (fünftausend Euro) auf "
    "6.000,00 € (sechstausend Euro) angepasst."
)
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY

run = p.add_run("\n\n")

run = p.add_run(
    "Alle übrigen Regelungen des bestehenden Geschäftsführer-Anstellungsvertrages "
    "bleiben unverändert."
)
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY

# ══════════════════════════════════════════════════════════════
# BEGRÜNDUNG
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Begründung")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

add_body(
    "Die Anpassung der Geschäftsführervergütung erfolgt aufgrund des gestiegenen "
    "Geschäftsumfangs und der zunehmenden Verantwortung des Geschäftsführers. "
    "Die neue Vergütung in Höhe von 72.000,00 € brutto p.a. liegt im Rahmen "
    "der marktüblichen Vergütung für Geschäftsführer vergleichbarer Unternehmen "
    "(IT-Dienstleistung / Beratung, Einzelgesellschafter)."
)

# ══════════════════════════════════════════════════════════════
# VERGÜTUNGSÜBERSICHT
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Vergütungsübersicht")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

# Styled table
table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

headers = ["", "Betrag"]
data = [
    ("Bisheriges Bruttogehalt (monatlich)", "5.000,00 €"),
    ("Neues Bruttogehalt (monatlich)", "6.000,00 €"),
]

# Header row
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_shading(cell, TEAL_HEX)
    set_cell_text(cell, header, bold=True, color=WHITE, size=Pt(10))

# Data rows
for r, (label, value) in enumerate(data):
    cell_label = table.rows[r + 1].cells[0]
    cell_value = table.rows[r + 1].cells[1]

    if r == 1:  # Highlight the new salary
        set_cell_shading(cell_label, "E6F7F7")
        set_cell_shading(cell_value, "E6F7F7")
        set_cell_text(cell_label, label, bold=True, color=TEAL, size=Pt(10))
        set_cell_text(cell_value, value, bold=True, color=TEAL, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
    else:
        set_cell_text(cell_label, label, size=Pt(10))
        set_cell_text(cell_value, value, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.RIGHT)

# Column widths
for row in table.rows:
    row.cells[0].width = Cm(10)
    row.cells[1].width = Cm(5)

# Effective date note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Wirksamkeit: ab 01. Februar 2026")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run.font.color.rgb = TEAL
run.font.bold = True

# ══════════════════════════════════════════════════════════════
# ABSTIMMUNGSERGEBNIS
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Abstimmungsergebnis")
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = TEAL
run.font.bold = True

add_body(
    "Der vorstehende Beschluss wurde vom alleinigen Gesellschafter "
    "einstimmig gefasst."
)

add_body(
    "Es wird festgestellt, dass der Beschluss ordnungsgemäß zustande "
    "gekommen ist."
)

# ══════════════════════════════════════════════════════════════
# SIGNATURE
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(4)

# Date and place
add_body("Frankfurt am Main, den 08. Februar 2026")

# Signature line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(50)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("_" * 40)
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Fayssal El Mofatiche")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.color.rgb = DARK_GRAY
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("Alleiniger Gesellschafter und Geschäftsführer")
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY

p = doc.add_paragraph()
run = p.add_run("Flowistic GmbH")
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = TEAL

# ══════════════════════════════════════════════════════════════
# FOOTER NOTE
# ══════════════════════════════════════════════════════════════

add_teal_line()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run(
    "Dieses Dokument ist zu den Geschäftsunterlagen der Gesellschaft zu nehmen."
)
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = DARK_GRAY
run.font.italic = True

# ── Save ──
os.makedirs(OUTPUT_DIR, exist_ok=True)
output_path = os.path.join(OUTPUT_DIR, "Gesellschafterbeschluss_Gehaltsanpassung_2026.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
