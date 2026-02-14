"""Bio command — generate a condensed one-pager bio."""

import os
from pathlib import Path
from typing import Annotated, Optional

import yaml
import typer
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .utils import (
    BASE_DIR, DEFAULT_YAML, OUTPUT_DIR, ASSETS_DIR,
    NAVY, NAVY_HEX, WHITE, SLATE, DARK_TEXT, TEAL, TEAL_HEX,
    LABELS, load_cv, open_file, convert_to_pdf, parse_start_date,
)
from .translate import translate_cv
from .docx_builder import add_para, add_run_to_para, set_cell_shading, remove_table_borders, remove_cell_borders, set_cell_width, set_cell_margins


def select_bio_content(cv: dict) -> dict:
    """Use Claude API to select and condense CV content for a one-pager bio."""
    from .utils import get_anthropic_client
    client = get_anthropic_client()

    cv_yaml = yaml.dump(cv, allow_unicode=True, default_flow_style=False, sort_keys=False)

    typer.echo("Generating bio content via Claude API...")
    message = client.messages.create(
        model="claude-sonnet-4-5-20250929",
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": (
                    "You are a professional CV consultant. Given a full CV in YAML, produce a condensed "
                    "one-pager bio version.\n\n"
                    "Return ONLY valid YAML with this exact structure — no explanation, no code fences:\n\n"
                    "name: <full name>\n"
                    "title: <professional title>\n"
                    "photo: <photo path from original>\n"
                    "contact:\n"
                    "  email: <email>\n"
                    "  phone: <phone>\n"
                    "  address: <city, country>\n"
                    "bio_summary: <3-4 sentences in third person summarizing the person's career, expertise, and value proposition>\n"
                    "career_highlights:\n"
                    "  - <highlight 1>\n"
                    "  - <highlight 2>\n"
                    "  - <highlight 3>\n"
                    "  - <highlight 4>\n"
                    "  - <highlight 5>\n"
                    "current_roles:\n"
                    "  - title: <title>\n"
                    "    org: <org>\n"
                    "    period: <start — end>\n"
                    "  ...(2-3 most recent roles)\n"
                    "education:\n"
                    "  - <degree, institution>\n"
                    "  ...\n"
                    "skills_summary: <comma-separated list of 10-15 key skills>\n"
                    "links:\n"
                    "  - label: <label>\n"
                    "    url: <url>\n"
                    "  ...\n\n"
                    "Rules:\n"
                    "- Do NOT invent content. Only select and condense from the original CV.\n"
                    "- Career highlights should be the most impressive, quantifiable achievements.\n"
                    "- Current roles = 2-3 most recent experience entries.\n"
                    "- Skills summary = most important skills across all categories.\n"
                    "- Bio summary should be written in third person, professional tone.\n\n"
                    f"CV YAML:\n{cv_yaml}"
                ),
            }
        ],
    )

    bio_yaml = message.content[0].text.strip()
    if bio_yaml.startswith("```"):
        bio_yaml = "\n".join(bio_yaml.split("\n")[1:])
    if bio_yaml.endswith("```"):
        bio_yaml = "\n".join(bio_yaml.split("\n")[:-1])

    return yaml.safe_load(bio_yaml)


def select_bio_content_deterministic(cv: dict) -> dict:
    """Deterministic fallback: select bio content without using the API."""
    # Sort experience by date, take top 3
    entries = sorted(
        cv.get("experience", []),
        key=lambda e: parse_start_date(e.get("start", "")),
        reverse=True,
    )
    current_roles = [
        {"title": e["title"], "org": e.get("org", ""), "period": f"{e['start']} — {e['end']}"}
        for e in entries[:3]
    ]

    # Collect top bullets as highlights
    highlights = []
    for e in entries[:4]:
        for b in e.get("bullets", [])[:2]:
            highlights.append(b)
            if len(highlights) >= 5:
                break
        if len(highlights) >= 5:
            break

    # Skills summary
    skills = cv.get("skills", {})
    all_skills = skills.get("leadership", []) + skills.get("technical", [])
    skills_summary = ", ".join(all_skills[:15])

    # Education
    education = [f"{e['degree']}, {e['institution']}" for e in cv.get("education", [])]

    # Profile as bio summary
    bio_summary = cv.get("profile", "").strip()
    # Take first 3 sentences
    sentences = bio_summary.split(". ")
    bio_summary = ". ".join(sentences[:4])
    if not bio_summary.endswith("."):
        bio_summary += "."

    return {
        "name": cv["name"],
        "title": cv["title"],
        "photo": cv.get("photo", ""),
        "contact": {
            "email": cv["contact"]["email"],
            "phone": cv["contact"]["phone"],
            "address": cv["contact"]["address"],
        },
        "bio_summary": bio_summary,
        "career_highlights": highlights,
        "current_roles": current_roles,
        "education": education,
        "skills_summary": skills_summary,
        "links": cv.get("links", []),
    }


def build_bio_docx(bio_data: dict, lang: str) -> Path:
    """Build a single-column one-pager bio document."""
    doc = Document()

    # Page setup — wider margins for single column
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Arial Narrow"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_TEXT
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── Header: photo + name + title + contact in a table ──
    header_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(header_table)

    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)

    # Set widths
    set_cell_width(left_cell, 4.0)
    set_cell_width(right_cell, 13.0)

    # Photo in left cell
    photo_path = BASE_DIR / bio_data.get("photo", "") if bio_data.get("photo") else None
    if photo_path and photo_path.exists():
        p_photo = left_cell.paragraphs[0]
        p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_photo.add_run()
        run.add_picture(str(photo_path), width=Cm(3.0))

    # Name + title + contact in right cell
    p_name = right_cell.paragraphs[0]
    run_name = p_name.add_run(bio_data["name"])
    run_name.bold = True
    run_name.font.size = Pt(18)
    run_name.font.color.rgb = NAVY
    run_name.font.name = "Arial Narrow"

    p_title = right_cell.add_paragraph()
    run_title = p_title.add_run(bio_data["title"])
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = TEAL
    run_title.font.name = "Arial Narrow"
    run_title.italic = True
    p_title.paragraph_format.space_after = Pt(6)

    # Contact line
    contact = bio_data.get("contact", {})
    contact_parts = [v for v in [contact.get("email"), contact.get("phone"), contact.get("address")] if v]
    if contact_parts:
        p_contact = right_cell.add_paragraph()
        run_c = p_contact.add_run("  |  ".join(contact_parts))
        run_c.font.size = Pt(8)
        run_c.font.color.rgb = SLATE
        run_c.font.name = "Arial Narrow"

    # Links
    links = bio_data.get("links", [])
    if links:
        p_links = right_cell.add_paragraph()
        link_labels = [l["label"] for l in links]
        run_l = p_links.add_run("  |  ".join(link_labels))
        run_l.font.size = Pt(8)
        run_l.font.color.rgb = TEAL
        run_l.font.name = "Arial Narrow"

    # ── Horizontal divider ──
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(8)
    p_div.paragraph_format.space_after = Pt(8)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{TEAL_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Bio summary ──
    def add_section_title(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        run.font.name = "Arial Narrow"

    add_section_title("Professional Summary")
    p_bio = doc.add_paragraph()
    run_bio = p_bio.add_run(bio_data.get("bio_summary", ""))
    run_bio.font.size = Pt(10)
    run_bio.font.color.rgb = DARK_TEXT
    run_bio.font.name = "Calibri"
    p_bio.paragraph_format.space_after = Pt(6)

    # ── Career highlights ──
    highlights = bio_data.get("career_highlights", [])
    if highlights:
        add_section_title("Key Achievements")
        for h in highlights:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.text = ""
            run = p.add_run(h)
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_TEXT
            run.font.name = "Calibri"

    # ── Current roles ──
    roles = bio_data.get("current_roles", [])
    if roles:
        add_section_title("Recent Experience")
        for role in roles:
            p_role = doc.add_paragraph()
            p_role.paragraph_format.space_before = Pt(4)
            p_role.paragraph_format.space_after = Pt(1)
            run_t = p_role.add_run(f"{role['title']} — {role.get('org', '')}")
            run_t.bold = True
            run_t.font.size = Pt(9)
            run_t.font.color.rgb = NAVY
            run_t.font.name = "Arial Narrow"

            p_period = doc.add_paragraph()
            run_p = p_period.add_run(role.get("period", ""))
            run_p.font.size = Pt(8)
            run_p.font.color.rgb = SLATE
            run_p.font.name = "Arial Narrow"

    # ── Education ──
    education = bio_data.get("education", [])
    if education:
        add_section_title("Education")
        for edu in education:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_after = Pt(2)
            edu_text = edu if isinstance(edu, str) else f"{edu.get('degree', '')}, {edu.get('institution', '')}"
            run_edu = p_edu.add_run(edu_text)
            run_edu.font.size = Pt(9)
            run_edu.font.color.rgb = DARK_TEXT
            run_edu.font.name = "Arial Narrow"

    # ── Skills summary ──
    skills_summary = bio_data.get("skills_summary", "")
    if skills_summary:
        add_section_title("Core Competencies")
        p_skills = doc.add_paragraph()
        run_skills = p_skills.add_run(skills_summary)
        run_skills.font.size = Pt(9)
        run_skills.font.color.rgb = DARK_TEXT
        run_skills.font.name = "Calibri"

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"Fayssal_El_Mofatiche_Bio_{lang.upper()}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    return output_path


def bio(
    lang: Annotated[Optional[str], typer.Option(help="Output language (en or de). Default: en.")] = "en",
    source: Annotated[Path, typer.Option(help="Path to source YAML.")] = DEFAULT_YAML,
    pdf: Annotated[bool, typer.Option("--pdf", help="Also generate PDF.")] = False,
    open: Annotated[bool, typer.Option("--open/--no-open", help="Open the generated file.")] = True,
):
    """Generate a condensed one-pager bio document."""
    cv_en = load_cv(source)

    # Try API, fall back to deterministic
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if api_key:
        bio_data = select_bio_content(cv_en)
    else:
        typer.echo("ANTHROPIC_API_KEY not set — using deterministic bio selection.")
        bio_data = select_bio_content_deterministic(cv_en)

    # Translate bio data if needed
    if lang == "de":
        bio_data = translate_cv(bio_data, retranslate=True)

    output_path = build_bio_docx(bio_data, lang)
    typer.echo(f"Generated: {output_path}")

    outputs = [output_path]
    if pdf:
        pdf_path = convert_to_pdf(output_path)
        typer.echo(f"Generated: {pdf_path}")
        outputs.append(pdf_path)

    if open:
        for p in outputs:
            open_file(p)
