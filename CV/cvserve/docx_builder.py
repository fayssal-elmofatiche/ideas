"""All docx building functions for the two-column CV layout."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .utils import (
    BASE_DIR, ASSETS_DIR, OUTPUT_DIR,
    NAVY, NAVY_HEX, WHITE, SLATE, DARK_TEXT, TEAL, TEAL_HEX,
    SIDEBAR_WIDTH_CM, MAIN_WIDTH_CM,
    PAGE_TOP_MARGIN, PAGE_BOTTOM_MARGIN, PAGE_LEFT_MARGIN, PAGE_RIGHT_MARGIN,
    SECTION_ICONS, LABELS, parse_start_date,
)


# ── Low-level helpers ──

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def remove_cell_borders(cell):
    """Remove all borders from a table cell."""
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0"/>'
        '</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)


def set_cell_width(cell, width_cm):
    """Set a cell's width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>'
    )
    tcPr.append(tcW)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in cm."""
    def to_twips(cm_val):
        return int(cm_val * 567)
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{to_twips(top)}" w:type="dxa"/>'
        f'  <w:left w:w="{to_twips(left)}" w:type="dxa"/>'
        f'  <w:bottom w:w="{to_twips(bottom)}" w:type="dxa"/>'
        f'  <w:right w:w="{to_twips(right)}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    cell._tc.get_or_add_tcPr().append(margins)


def add_para(cell, text="", bold=False, italic=False, size=Pt(9), color=WHITE,
             font_name="Arial Narrow", align=None, space_before=Pt(0), space_after=Pt(0)):
    """Add a formatted paragraph to a cell."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = size
        run.font.color.rgb = color
        run.font.name = font_name
    return p


def add_run_to_para(p, text, bold=False, italic=False, size=Pt(9), color=DARK_TEXT,
                    font_name="Arial Narrow"):
    """Add a formatted run to an existing paragraph."""
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.color.rgb = color
    run.font.name = font_name
    return run


# ── Sidebar builders ──

def build_sidebar_header(cell, cv, lang="en"):
    """Build the header portion of the sidebar: photo + name."""
    photo_path = BASE_DIR / cv.get("photo", "") if cv.get("photo") else None

    if photo_path and photo_path.exists():
        p_photo = cell.paragraphs[0]  # Use existing first paragraph
        p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_photo.paragraph_format.space_after = Pt(8)
        run = p_photo.add_run()
        run.add_picture(str(photo_path), width=Cm(2.8))

    # Name — split after first name so last name + credential stay on one line
    name = cv["name"]
    parts = name.split(" ", 1)
    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(4)
    if len(parts) == 2:
        run1 = p_name.add_run(parts[0] + "\n")
        run1.bold = True
        run1.font.size = Pt(13)
        run1.font.color.rgb = WHITE
        run1.font.name = "Arial Narrow"
        run2 = p_name.add_run(parts[1])
        run2.bold = True
        run2.font.size = Pt(13)
        run2.font.color.rgb = WHITE
        run2.font.name = "Arial Narrow"
    else:
        run1 = p_name.add_run(name)
        run1.bold = True
        run1.font.size = Pt(13)
        run1.font.color.rgb = WHITE
        run1.font.name = "Arial Narrow"

    # Title
    add_para(cell, cv["title"], size=Pt(8), color=SLATE, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))


def build_sidebar_contact(cell, cv, lang="en"):
    """Build the contact details section in the sidebar."""
    L = LABELS[lang]
    # Section heading
    add_para(cell, L["details"], bold=True, size=Pt(10), color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(6), space_after=Pt(6))

    contact = cv["contact"]
    for text in [contact["address"], contact["phone"], contact["email"]]:
        add_para(cell, text, size=Pt(8), color=WHITE,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

    # Nationality
    add_para(cell, "", space_before=Pt(4))
    add_para(cell, L["nationality"], bold=True, size=Pt(9), color=SLATE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_para(cell, contact["nationality"], size=Pt(8), color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))


def build_sidebar_links(cell, cv, lang="en"):
    """Build the links section in the sidebar."""
    L = LABELS[lang]
    if not cv.get("links"):
        return
    add_para(cell, L["links"], bold=True, size=Pt(10), color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(8), space_after=Pt(6))

    for link in cv["links"]:
        add_para(cell, link["label"], size=Pt(8), color=TEAL,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))


def build_sidebar_skills(cell, cv, lang="en"):
    """Build the skills section in the sidebar."""
    L = LABELS[lang]
    skills = cv.get("skills", {})
    if not skills:
        return

    add_para(cell, L["skills"], bold=True, size=Pt(10), color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(10), space_after=Pt(6))

    # Leadership heading
    if skills.get("leadership"):
        add_para(cell, L["leadership_skills"], bold=True, size=Pt(8), color=SLATE,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
        for skill in skills["leadership"]:
            add_para(cell, skill, size=Pt(8), color=WHITE,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

    # Technical skills
    if skills.get("technical"):
        add_para(cell, "", space_before=Pt(4))
        for skill in skills["technical"]:
            add_para(cell, skill, size=Pt(8), color=WHITE,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))


def build_sidebar_languages(cell, cv, lang="en"):
    """Build the languages section in the sidebar."""
    L = LABELS[lang]
    languages = cv.get("skills", {}).get("languages", [])
    if not languages:
        return

    level_map_en = {"native": "Native", "fluent": "Fluent", "professional": "Professional", "basic": "Basic"}
    level_map_de = {"native": "Muttersprache", "fluent": "Fließend", "professional": "Verhandlungssicher", "basic": "Grundkenntnisse"}
    level_map = level_map_de if lang == "de" else level_map_en

    add_para(cell, L["languages"], bold=True, size=Pt(10), color=WHITE,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=Pt(10), space_after=Pt(6))

    for lang in languages:
        level = level_map.get(lang["level"], lang["level"])
        add_para(cell, f"{lang['name']} ({level})", size=Pt(8), color=WHITE,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))


# ── Main content builders ──

def add_section_heading(cell, title, icon_key=None):
    """Add a section heading row with icon to main content."""
    icon_file = SECTION_ICONS.get(icon_key or title)
    icon_path = ASSETS_DIR / icon_file if icon_file else None

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

    if icon_path and icon_path.exists():
        run_icon = p.add_run()
        run_icon.add_picture(str(icon_path), height=Pt(14))
        p.add_run("  ")

    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = NAVY
    run_title.font.name = "Arial Narrow"

    # Horizontal line
    p_line = cell.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(4)
    pPr = p_line._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{TEAL_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_labeled_line(cell, label, value, label_color=TEAL, value_color=SLATE):
    """Add a 'Label: value' line to main content."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_run_to_para(p, f"{label}: ", bold=True, size=Pt(8), color=label_color)
    add_run_to_para(p, value, size=Pt(8), color=value_color)


def build_main_profile(cell, cv, lang="en"):
    """Build the Profile section in main content."""
    L = LABELS[lang]
    add_section_heading(cell, L["profile"], icon_key="Profile")
    add_para(cell, cv["profile"].strip(), size=Pt(9), color=DARK_TEXT,
             font_name="Calibri", space_after=Pt(4))

    # Testimonials
    if cv.get("testimonials"):
        add_para(cell, L["testimonials_heading"], bold=True, size=Pt(9), color=NAVY,
                 space_before=Pt(8), space_after=Pt(4))
        for t in cv["testimonials"]:
            # Quote
            if t.get("quote"):
                add_para(cell, t["quote"], italic=True, size=Pt(9), color=SLATE,
                         font_name="Calibri", space_after=Pt(4))
            # Attribution
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            add_run_to_para(p, t["name"], bold=True, size=Pt(9), color=NAVY)
            add_run_to_para(p, f"\n{t['role']}", size=Pt(8), color=SLATE)
            add_run_to_para(p, f"\n{t['org']}", size=Pt(8), color=SLATE)


def build_main_experience(cell, cv, lang="en"):
    """Build the Experience section in main content."""
    L = LABELS[lang]
    add_section_heading(cell, L["experience"], icon_key="Project / Employment History")

    entries = sorted(cv.get("experience", []), key=lambda e: parse_start_date(e.get("start", "")), reverse=True)
    for entry in entries:
        # Title + org
        title_text = entry["title"]
        if entry.get("org"):
            title_text += f" — {entry['org']}"

        p_title = cell.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        p_title.paragraph_format.space_after = Pt(1)
        add_run_to_para(p_title, title_text, bold=True, size=Pt(9), color=NAVY)

        # Date
        add_para(cell, f"{entry['start']} — {entry['end']}", size=Pt(8), color=SLATE,
                 space_after=Pt(2))

        # Description
        if entry.get("description"):
            add_para(cell, entry["description"], italic=True, size=Pt(9), color=SLATE,
                     font_name="Calibri", space_after=Pt(2))

        # Bullets
        for bullet in entry.get("bullets", []):
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.text = ""
            add_run_to_para(p, bullet, size=Pt(8), color=DARK_TEXT, font_name="Calibri")

        # Metadata lines
        meta_fields = [
            ("soft_skills", "Soft Skills"),
            ("investment_skills", "Investment Skills"),
            ("telecom_skills", "Telecommunication Skills"),
            ("tech_stack", "Tech Stack"),
            ("frontend_tech_stack", "Frontend Tech Stack"),
            ("backend_tech_stack", "Backend Tech Stack"),
            ("devops", "DevOps / MLOps"),
            ("tools_platforms", "Tools & Platforms"),
            ("focus_areas", "Focus Expertise Areas"),
            ("project_methodology", "Project Methodology"),
            ("architecture", "Architecture"),
            ("project_size", "Project Size & Setup"),
        ]
        for key, label in meta_fields:
            val = entry.get(key)
            if val and (isinstance(val, str) or (isinstance(val, list) and len(val) > 0)):
                text = val if isinstance(val, str) else ", ".join(val)
                add_labeled_line(cell, label, text)


def build_main_education(cell, cv, lang="en"):
    """Build the Education section in main content."""
    L = LABELS[lang]
    add_section_heading(cell, L["education"], icon_key="Education")

    for entry in cv.get("education", []):
        p_title = cell.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(1)
        add_run_to_para(p_title, f"{entry['degree']}, {entry['institution']}",
                        bold=True, size=Pt(9), color=NAVY)

        add_para(cell, f"{entry['start']} — {entry['end']}", size=Pt(8), color=SLATE,
                 space_after=Pt(2))

        if entry.get("description"):
            add_para(cell, entry["description"], size=Pt(9), color=DARK_TEXT,
                     font_name="Calibri", space_after=Pt(2))

        if entry.get("details"):
            add_para(cell, entry["details"], size=Pt(8), color=SLATE,
                     font_name="Calibri", space_after=Pt(2))


def build_main_volunteering(cell, cv, lang="en"):
    """Build the Volunteering section in main content."""
    L = LABELS[lang]
    if not cv.get("volunteering"):
        return
    add_section_heading(cell, L["volunteering"], icon_key="Volunteering")

    for entry in cv["volunteering"]:
        p_title = cell.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(1)
        add_run_to_para(p_title, f"{entry['title']} — {entry['org']}",
                        bold=True, size=Pt(9), color=NAVY)

        add_para(cell, f"{entry['start']} — {entry['end']}", size=Pt(8), color=SLATE,
                 space_after=Pt(2))

        if entry.get("description"):
            add_para(cell, entry["description"], size=Pt(9), color=DARK_TEXT,
                     font_name="Calibri", space_after=Pt(2))


def build_main_references(cell, cv, lang="en"):
    """Build the References section in main content."""
    L = LABELS[lang]
    if not cv.get("references"):
        return
    add_section_heading(cell, L["references"], icon_key="References")
    add_para(cell, cv["references"], size=Pt(9), color=DARK_TEXT,
             font_name="Calibri", space_after=Pt(4))


def build_main_certifications(cell, cv, lang="en"):
    """Build the Certifications section in main content."""
    L = LABELS[lang]
    if not cv.get("certifications"):
        return
    add_section_heading(cell, L["certifications"], icon_key="Certifications")

    for entry in cv["certifications"]:
        title_text = entry["name"]
        if entry.get("org"):
            title_text += f", {entry['org']}"

        p_title = cell.add_paragraph()
        p_title.paragraph_format.space_before = Pt(6)
        p_title.paragraph_format.space_after = Pt(1)
        add_run_to_para(p_title, title_text, bold=True, size=Pt(9), color=NAVY)

        add_para(cell, f"{entry['start']} — {entry['end']}", size=Pt(8), color=SLATE,
                 space_after=Pt(2))

        if entry.get("description"):
            add_para(cell, entry["description"], size=Pt(9), color=DARK_TEXT,
                     font_name="Calibri", space_after=Pt(2))


def build_main_publications(cell, cv, lang="en"):
    """Build the Publications section in main content."""
    L = LABELS[lang]
    if not cv.get("publications"):
        return
    add_section_heading(cell, L["publications"], icon_key="Publications")

    for pub in cv["publications"]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        add_run_to_para(p, pub["title"], bold=True, size=Pt(9), color=NAVY)
        add_run_to_para(p, f"\n{pub['year']}: {pub['venue']}", size=Pt(8), color=SLATE)


# ── Document assembly ──

def build_docx(cv: dict, lang: str) -> Path:
    """Build a complete Word document from CV data."""
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = PAGE_TOP_MARGIN
        section.bottom_margin = PAGE_BOTTOM_MARGIN
        section.left_margin = PAGE_LEFT_MARGIN
        section.right_margin = PAGE_RIGHT_MARGIN

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Arial Narrow"
    style.font.size = Pt(9)
    style.font.color.rgb = DARK_TEXT
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Outer table: 1 row, 2 columns (sidebar | main)
    table = doc.add_table(rows=1, cols=2)
    remove_table_borders(table)

    # Set column widths via grid
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn('w:gridCol'))
        if len(gridCols) >= 2:
            gridCols[0].set(qn('w:w'), str(int(SIDEBAR_WIDTH_CM * 567)))
            gridCols[1].set(qn('w:w'), str(int(MAIN_WIDTH_CM * 567)))

    sidebar = table.cell(0, 0)
    main = table.cell(0, 1)

    # Style sidebar cell
    set_cell_shading(sidebar, NAVY_HEX)
    remove_cell_borders(sidebar)
    set_cell_width(sidebar, SIDEBAR_WIDTH_CM)
    set_cell_margins(sidebar, top=0.4, bottom=0.5, left=0.3, right=0.3)

    # Style main cell
    remove_cell_borders(main)
    set_cell_width(main, MAIN_WIDTH_CM)
    set_cell_margins(main, top=0.3, bottom=0.5, left=0.5, right=0.3)

    # ── Build sidebar ──
    build_sidebar_header(sidebar, cv, lang)
    build_sidebar_contact(sidebar, cv, lang)
    build_sidebar_links(sidebar, cv, lang)
    build_sidebar_skills(sidebar, cv, lang)
    build_sidebar_languages(sidebar, cv, lang)

    # ── Build main content ──
    # Clear the default empty paragraph
    main.paragraphs[0].text = ""

    build_main_profile(main, cv, lang)
    build_main_experience(main, cv, lang)
    build_main_education(main, cv, lang)
    build_main_volunteering(main, cv, lang)
    build_main_references(main, cv, lang)
    build_main_certifications(main, cv, lang)
    build_main_publications(main, cv, lang)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"Fayssal_El_Mofatiche_CV_{lang.upper()}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    return output_path
